"""Общие фикстуры для тестов."""

import os
from collections.abc import AsyncGenerator, Generator
from pathlib import Path
from unittest.mock import MagicMock

import pytest
import pytest_asyncio
from sqlalchemy import create_engine, text
from sqlalchemy.ext.asyncio import (
    AsyncSession,
    async_sessionmaker,
    create_async_engine,
)
from sqlalchemy.orm import Session, sessionmaker

# --- Моки для тяжёлых зависимостей ---
import sys

for name in ("torch", "transformers", "qwen_vl_utils"):
    sys.modules.setdefault(name, MagicMock())

# --- PostgreSQL ---

TEST_DATABASE_URL = os.getenv(
    "TEST_DATABASE_URL",
    "postgresql+asyncpg://cv_user:changeme@localhost:5432/cv_analyzer",
)

MIGRATION_SQL = Path("migrations/001_init.sql").read_text()

DROP_ALL_SQL = """\
DROP TRIGGER IF EXISTS trg_tasks_updated_at ON tasks;
DROP FUNCTION IF EXISTS update_updated_at() CASCADE;
DROP TABLE IF EXISTS additional, skills, experience, education,
                  personal_data, resumes, tasks CASCADE;
"""

TRUNCATE_SQL = """\
TRUNCATE TABLE additional, skills, experience, education,
               personal_data, resumes, tasks CASCADE;
"""


def _apply_migrations_via_sync_engine(url: str):
    """Создаёт временный sync-engine, накатывает миграции, закрывает."""
    sync_url = url.replace("+asyncpg", "")
    engine = create_engine(sync_url, echo=False, isolation_level="AUTOCOMMIT")
    try:
        with engine.begin() as conn:
            conn.execute(text(DROP_ALL_SQL))
            conn.execute(text(MIGRATION_SQL))
    finally:
        engine.dispose()


# ---------------------------------------------------------------------------
# Session-scoped: миграции (один раз за весь прогон)
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def _pg_ready():
    """Накатывает миграции до тестов, дропает после."""
    try:
        _apply_migrations_via_sync_engine(TEST_DATABASE_URL)
    except Exception as exc:
        pytest.skip(f"PostgreSQL недоступна: {exc}")
    yield
    _apply_migrations_via_sync_engine(TEST_DATABASE_URL)


# ---------------------------------------------------------------------------
# Async fixtures (API CRUD, endpoints) — function-scoped
# ---------------------------------------------------------------------------


@pytest_asyncio.fixture
async def async_engine(_pg_ready):
    """Function-scoped: свежий async engine для каждого теста."""
    engine = create_async_engine(TEST_DATABASE_URL, echo=False)
    async with engine.begin() as conn:
        await conn.execute(text(TRUNCATE_SQL))
    yield engine
    await engine.dispose()


@pytest_asyncio.fixture
async def async_session(async_engine) -> AsyncGenerator[AsyncSession, None]:
    """Function-scoped: чистая БД + async-сессия."""
    factory = async_sessionmaker(
        async_engine, class_=AsyncSession, expire_on_commit=False
    )
    async with factory() as session:
        yield session


# ---------------------------------------------------------------------------
# Sync fixtures (Worker CRUD)
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def sync_engine(_pg_ready):
    """Session-scoped: sync PostgreSQL engine."""
    sync_url = TEST_DATABASE_URL.replace("+asyncpg", "")
    engine = create_engine(sync_url, echo=False)
    yield engine
    engine.dispose()


@pytest.fixture
def sync_session(sync_engine) -> Generator[Session, None, None]:
    """Function-scoped: чистая БД + sync-сессия."""
    with sync_engine.begin() as conn:
        conn.execute(text(TRUNCATE_SQL))
    factory = sessionmaker(
        bind=sync_engine, class_=Session, expire_on_commit=False
    )
    session = factory()
    yield session
    session.close()


# ---------------------------------------------------------------------------
# Тестовые файлы
# ---------------------------------------------------------------------------


def _build_pdf(text: str) -> bytes:
    """Генерирует минимальный валидный PDF с заданным текстом."""
    content_line = f"BT /F1 12 Tf 100 700 Td ({text}) Tj ET".encode("latin-1")
    objects = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        4: f"<< /Length {len(content_line)} >>\nstream\n".encode("latin-1")
        + content_line
        + b"\nendstream",
        5: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }

    pdf = b"%PDF-1.4\n"
    offsets = {}
    for num in sorted(objects):
        offsets[num] = len(pdf)
        pdf += f"{num} 0 obj\n".encode() + objects[num] + b"\nendobj\n"

    xref_offset = len(pdf)
    pdf += b"xref\n0 6\n"
    pdf += b"0000000000 65535 f \n"
    for num in sorted(offsets):
        pdf += f"{offsets[num]:010d} 00000 n \n".encode()

    pdf += f"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode()
    return pdf


@pytest.fixture
def sample_pdf(tmp_path):
    """Минимальный PDF с текстом."""
    path = tmp_path / "sample.pdf"
    path.write_bytes(_build_pdf("Ivan Ivanov Python Developer"))
    return path


@pytest.fixture
def empty_pdf(tmp_path):
    """PDF без текста (пустая страница)."""
    objects = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] >>",
    }
    pdf = b"%PDF-1.4\n"
    offsets = {}
    for num in sorted(objects):
        offsets[num] = len(pdf)
        pdf += f"{num} 0 obj\n".encode() + objects[num] + b"\nendobj\n"

    xref_offset = len(pdf)
    pdf += b"xref\n0 4\n"
    pdf += b"0000000000 65535 f \n"
    for num in sorted(offsets):
        pdf += f"{offsets[num]:010d} 00000 n \n".encode()
    pdf += f"trailer\n<< /Size 4 /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode()

    path = tmp_path / "empty.pdf"
    path.write_bytes(pdf)
    return path


@pytest.fixture
def sample_docx(tmp_path):
    """Минимальный DOCX с текстом."""
    from docx import Document

    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Ivan Ivanov")
    doc.add_paragraph("Python Developer")
    doc.add_paragraph("Experience: 5 years")
    doc.save(str(path))
    return path


@pytest.fixture
def sample_odt(tmp_path):
    """Минимальный ODT с текстом."""
    from odf.opendocument import OpenDocumentText
    from odf.text import P

    path = tmp_path / "sample.odt"
    doc = OpenDocumentText()
    for t in ("Ivan Ivanov", "Python Developer"):
        doc.text.addElement(P(text=t))
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Тестовые изображения
# ---------------------------------------------------------------------------


def _build_jpeg(width: int = 200, height: int = 300, mode: str = "RGB") -> bytes:
    """Генерирует минимальный JPEG с белым фоном."""
    from io import BytesIO

    from PIL import Image

    img = Image.new(mode, (width, height), color=(255, 255, 255))
    buf = BytesIO()
    img.save(buf, format="JPEG")
    return buf.getvalue()


def _build_png(width: int = 200, height: int = 300, mode: str = "RGBA") -> bytes:
    """Генерирует минимальный PNG с прозрачным фоном."""
    from io import BytesIO

    from PIL import Image

    img = Image.new(mode, (width, height), color=(0, 0, 0, 0))
    buf = BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


@pytest.fixture
def sample_jpeg(tmp_path):
    """Минимальный JPEG-файл (RGB, 200x300)."""
    path = tmp_path / "sample.jpeg"
    path.write_bytes(_build_jpeg())
    return path


@pytest.fixture
def sample_jpg(tmp_path):
    """Минимальный JPG-файл (расширение .jpg)."""
    path = tmp_path / "sample.jpg"
    path.write_bytes(_build_jpeg())
    return path


@pytest.fixture
def sample_png(tmp_path):
    """Минимальный PNG-файл (RGBA, 200x300)."""
    path = tmp_path / "sample.png"
    path.write_bytes(_build_png())
    return path


@pytest.fixture
def large_jpeg(tmp_path):
    """JPEG размером больше MAX_IMAGE_SIZE (1000x1200)."""
    path = tmp_path / "large.jpeg"
    path.write_bytes(_build_jpeg(width=1000, height=1200))
    return path


@pytest.fixture
def grayscale_jpeg(tmp_path):
    """JPEG в режиме Grayscale (не RGB)."""
    from io import BytesIO

    from PIL import Image

    img = Image.new("L", (200, 300), color=128)
    buf = BytesIO()
    img.save(buf, format="JPEG")
    path = tmp_path / "gray.jpeg"
    path.write_bytes(buf.getvalue())
    return path


@pytest.fixture
def corrupted_image(tmp_path):
    """Битый файл изображения (мусорные байты)."""
    path = tmp_path / "broken.jpeg"
    path.write_bytes(b"\xff\xd8\xff\xe0" + b"\x00" * 50)
    return path


@pytest.fixture
def varied_jpeg(tmp_path):
    """JPEG с разными значениями пикселей (для тестов contrast/brightness)."""
    from io import BytesIO

    from PIL import Image

    img = Image.new("RGB", (10, 10), color=(128, 64, 32))
    buf = BytesIO()
    img.save(buf, format="JPEG")
    path = tmp_path / "varied.jpeg"
    path.write_bytes(buf.getvalue())
    return path
