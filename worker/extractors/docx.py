"""Извлечение текста из DOCX-файлов через python-docx."""

import logging

from docx import Document

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> str:
    """Извлекает текст всех параграфов и таблиц DOCX."""
    doc = Document(file_path)
    text_parts = []

    # Извлекаем параграфы из основного текста
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Извлекаем текст из таблиц (очень часто в резюме)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    result = "\n".join(text_parts)
    logger.info("DOCX: извлечено %d параграфов из %s", len(text_parts), file_path)
    return result
